"""
Word Document Generator for Municipality Data
Creates Word documents with municipality equipment data
Supports multiple countries with localized labels
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import sys
import os

# Handle path for both script and PyInstaller .exe
if getattr(sys, 'frozen', False):
    sys.path.insert(0, sys._MEIPASS)
else:
    sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from config import (
    OUTPUT_WORD,
    SIMPLE_COLUMNS,
    DATA_DIR,
    POPULATION_THRESHOLD,
    MIN_RURAL_POPULATION,
    EQUIPMENT_DIVISOR_URBAN,
    EQUIPMENT_DIVISOR_RURAL,
    DEFAULT_URBAN_PERCENTAGE
)


class WordGenerator:
    """Generator for Word documents with municipality data"""

    def __init__(self, country_config=None):
        """
        Initialize generator with optional country configuration

        Args:
            country_config: Dictionary with country-specific settings
        """
        self.header_color = "4472C4"  # Blue

        # Set country-specific labels or use defaults
        if country_config and "labels" in country_config:
            labels = country_config["labels"]
            self.simple_columns = {
                'name': labels.get('municipality', SIMPLE_COLUMNS['name']),
                'total_equipos': labels.get('total_equipos', SIMPLE_COLUMNS['total_equipos']),
            }
            self.country_name = country_config.get("name", "España")
        else:
            self.simple_columns = SIMPLE_COLUMNS
            self.country_name = "España"

        # Generate country-specific file name
        country_suffix = self._get_country_suffix()
        self.output_word = os.path.join(DATA_DIR, f"municipios_{country_suffix}_clasificacion.docx")

    def _get_country_suffix(self):
        """Get lowercase country name for file naming"""
        country_map = {
            "España": "espana",
            "France": "france",
            "Italia": "italia",
            "Portugal": "portugal",
            "Deutschland": "deutschland",
        }
        return country_map.get(self.country_name, self.country_name.lower().replace(" ", "_"))

    def _ensure_data_dir(self):
        """Ensure data directory exists"""
        os.makedirs(DATA_DIR, exist_ok=True)

    def _set_cell_shading(self, cell, color):
        """Set background color for a table cell"""
        shading_elm = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{color}"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def _calculate_total_equipos(self, population):
        """
        Calculate total equipment based on population

        Logic:
        - Rural (< 2000): EQUIPOS = population / 50
        - Urban (>= 2000): EQUIPOS = (hab_urban / 300) + (2000 / 50)
          where hab_urban = population - 2000
        """
        if population is None or population == 0:
            return 0

        is_urban = population >= POPULATION_THRESHOLD

        if is_urban:
            hab_rural = MIN_RURAL_POPULATION  # Fixed 2000
            hab_urban = population - MIN_RURAL_POPULATION
            equipos_rural = hab_rural / EQUIPMENT_DIVISOR_RURAL
            equipos_urban = hab_urban / EQUIPMENT_DIVISOR_URBAN
            return round(equipos_urban + equipos_rural, 2)
        else:
            return round(population / EQUIPMENT_DIVISOR_RURAL, 2)

    def create_word_document(self, municipalities, output_path=None):
        """
        Create Word document with 2-column table (Municipality and TOTAL EQUIPOS)

        Args:
            municipalities: List of municipality dictionaries
            output_path: Output file path (optional, uses config default)

        Returns:
            Path to created file
        """
        self._ensure_data_dir()
        output_path = output_path or self.output_word

        # Create document
        doc = Document()

        # Add title
        title = doc.add_heading(f'Equipos de Municipios de {self.country_name}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add subtitle with count
        subtitle = doc.add_paragraph(f'Total de municipios: {len(municipalities)}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add space
        doc.add_paragraph()

        # Create table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = self.simple_columns['name']
        header_cells[1].text = self.simple_columns['total_equipos']

        # Style header
        for cell in header_cells:
            self._set_cell_shading(cell, self.header_color)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)

        # Add data rows
        for m in municipalities:
            row_cells = table.add_row().cells
            row_cells[0].text = m.get('name', '')

            # Calculate total equipos
            population = m.get('population') or 0
            total_equipos = self._calculate_total_equipos(population)
            row_cells[1].text = str(total_equipos)

            # Center the equipos column
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Set column widths
        for row in table.rows:
            row.cells[0].width = Inches(3)
            row.cells[1].width = Inches(2)

        # Save document
        doc.save(output_path)
        print(f"Word document saved: {output_path}")

        return output_path


def main():
    """Test the Word generator"""
    generator = WordGenerator()

    # Test data
    test_municipalities = [
        {'name': 'Tudela', 'population': 37008},
        {'name': 'Tafalla', 'population': 10582},
        {'name': 'Sartaguda', 'population': 1287},
        {'name': 'Sesma', 'population': 1149},
    ]

    generator.create_word_document(test_municipalities)
    print("Test Word document created successfully!")


if __name__ == "__main__":
    main()
